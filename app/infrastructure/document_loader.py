import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

from app.core.settings import settings
from app.domain.models import DocumentChunk


def _normalize_identifier(value: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9ก-๙]+", "-", value.strip()).strip("-").lower()
    return normalized or settings.DEFAULT_COLLECTION_ID


@dataclass(frozen=True)
class DiscoveredDocument:
    path: Path
    relative_path: str
    collection_id: str
    document_name: str
    file_mtime: float
    file_hash: str


class DocumentLoader:
    def __init__(self, target_path: str):
        self.target_path = Path(target_path)

    @staticmethod
    def _hash_file(file_path: Path) -> str:
        return hashlib.sha1(file_path.read_bytes()).hexdigest()

    def _resolve_collection_id(self, file_path: Path) -> str:
        base_dir = self.target_path if self.target_path.is_dir() else settings.resolved_doc_path

        try:
            relative_parent = file_path.parent.relative_to(base_dir)
        except ValueError:
            return settings.DEFAULT_COLLECTION_ID

        if not relative_parent.parts:
            return settings.DEFAULT_COLLECTION_ID

        return _normalize_identifier("__".join(relative_parent.parts))

    def _resolve_relative_path(self, file_path: Path) -> str:
        if file_path.is_relative_to(settings.resolved_doc_path):
            return file_path.relative_to(settings.resolved_doc_path).as_posix()
        return file_path.name

    def _build_chunk_id(
        self,
        collection_id: str,
        document: str,
        chapter: str | None,
        article: str | None,
        index: int,
        text: str,
    ) -> str:
        identity = "|".join([collection_id, document, chapter or "", article or "", str(index), text])
        return hashlib.sha1(identity.encode("utf-8")).hexdigest()

    def discover_documents(self) -> list[DiscoveredDocument]:
        files_to_process: list[Path] = []

        if self.target_path.is_file() and self.target_path.suffix == ".docx":
            files_to_process = [self.target_path]
        elif self.target_path.is_dir():
            files_to_process = sorted(self.target_path.rglob("*.docx"))

        discovered_documents: list[DiscoveredDocument] = []
        for file_path in files_to_process:
            stats = file_path.stat()
            discovered_documents.append(
                DiscoveredDocument(
                    path=file_path,
                    relative_path=self._resolve_relative_path(file_path),
                    collection_id=self._resolve_collection_id(file_path),
                    document_name=file_path.stem,
                    file_mtime=stats.st_mtime,
                    file_hash=self._hash_file(file_path),
                )
            )

        return discovered_documents

    def load_documents(self, documents: list[DiscoveredDocument] | None = None) -> dict[str, list[DocumentChunk]]:
        import docx

        discovered_documents = documents or self.discover_documents()
        chunks_by_source: dict[str, list[DocumentChunk]] = {}

        for document in discovered_documents:
            doc = docx.Document(str(document.path))
            current_chapter = ""
            current_article = "บททั่วไป/คำปรารภ"
            source_chunks: list[DocumentChunk] = []

            for index, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue

                match_chapter = re.match(r"^(หมวด\s*[๐-๙0-9]+\s*.*)", text)
                if match_chapter:
                    current_chapter = match_chapter.group(1).strip()
                    chunk_text = f"[{document.document_name}] {text}"
                    chapter = current_chapter
                    article = None
                else:
                    match_article = re.match(r"^((?:มาตรา|ข้อ)\s*[๐-๙0-9/]+)", text)
                    if match_article:
                        current_article = match_article.group(1).strip()
                        tag = f"[{document.document_name}]"
                        if current_chapter:
                            tag += f" [{current_chapter}]"
                        chunk_text = f"{tag} {text}"
                        chapter = current_chapter or None
                        article = current_article
                    else:
                        tag = f"[{document.document_name}]"
                        if current_chapter:
                            tag += f" [{current_chapter}]"
                        tag += f" [{current_article}]"
                        chunk_text = f"{tag} {text}"
                        chapter = current_chapter or None
                        article = current_article

                chunk_index = len(source_chunks)
                source_chunks.append(
                    DocumentChunk(
                        id=self._build_chunk_id(
                            document.collection_id,
                            document.document_name,
                            chapter,
                            article,
                            chunk_index,
                            chunk_text,
                        ),
                        text=chunk_text,
                        document=document.document_name,
                        collection_id=document.collection_id,
                        chapter=chapter,
                        article=article,
                        metadata={
                            "document": document.document_name,
                            "collection_id": document.collection_id,
                            "chapter": chapter or "",
                            "article": article or "",
                            "source_path": document.relative_path,
                            "source_mtime": document.file_mtime,
                            "source_hash": document.file_hash,
                            "embedding_model": settings.EMBED_MODEL_NAME,
                            "pipeline_version": settings.VECTOR_PIPELINE_VERSION,
                            "chunk_index": chunk_index,
                        },
                    )
                )

            chunks_by_source[document.relative_path] = source_chunks

        return chunks_by_source

    def load(self) -> list[DocumentChunk]:
        chunks: list[DocumentChunk] = []
        for source_chunks in self.load_documents().values():
            chunks.extend(source_chunks)
        return chunks
